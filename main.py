
from fastapi import FastAPI, UploadFile, File, Form
from fastapi.middleware.cors import CORSMiddleware
import pandas as pd
import io
import json
import numpy as np
import pingouin as pg
from scipy import stats
from docx import Document

app = FastAPI()

# --- CORS (Handshake with your website) ---
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

def run_ultimate_computation_engine(df, options):
    # Standardize column headers (remove leading/trailing spaces)
    df.columns = df.columns.str.strip()
    
    # 1. AUTO-CLEANING: Removes %, $, and commas to fix 'Protocol Failure'
    for col in df.columns:
        if df[col].dtype == 'object':
            df[col] = df[col].astype(str).str.replace(r'[%\$,]', '', regex=True)
            df[col] = pd.to_numeric(df[col], errors='coerce')

    # Get numeric columns for processing
    num_cols = df.select_dtypes(include=[np.number]).columns
    results = {}

    # 2. DESCRIPTIVE STATISTICS (Always included)
    results["descriptive"] = {col: {
        "n": int(df[col].count()), 
        "mean": float(df[col].mean()),
        "std": float(df[col].std()), 
        "min": float(df[col].min()),
        "max": float(df[col].max()), 
        "skew": float(df[col].skew()),
        "kurt": float(df[col].kurtosis())
    } for col in num_cols}

    # 3. NORMALITY TEST (Shapiro-Wilk)
    results["normality"] = {col: {"p_value": float(stats.shapiro(df[col].dropna())[1])} 
                            for col in num_cols if len(df[col].dropna()) > 3}

    # 4. INFERENTIAL: PINGOUIN T-TEST (Fuzzy Gender/Sex matching)
    try:
        g_col = next((c for c in df.columns if 'gender' in c.lower() or 'sex' in c.lower()), None)
        v_col = num_cols[0] if len(num_cols) > 0 else None
        if g_col and v_col:
            m = df[df[g_col].astype(str).str.lower().str.startswith('m', na=False)][v_col].dropna()
            f = df[df[g_col].astype(str).str.lower().str.startswith('f', na=False)][v_col].dropna()
            if len(m) > 1 and len(f) > 1:
                t_res = pg.ttest(m, f)
                results["t_test"] = {
                    "t_stat": float(t_res['T'].iloc[0]), 
                    "p_value": float(t_res['p-val'].iloc[0]),
                    "df": int(t_res['dof'].iloc[0]), 
                    "significant": bool(t_res['p-val'].iloc[0] < 0.05),
                    "cohens_d": float(t_res['cohen-d'].iloc[0])
                }
    except: pass

    # 5. INFERENTIAL: ONE-WAY ANOVA (Fuzzy Group/Class matching)
    try:
        grp_col = next((c for c in df.columns if 'group' in c.lower() or 'class' in c.lower()), None)
        if grp_col and v_col:
            unique_grps = df[grp_col].dropna().unique()
            if len(unique_grps) > 2:
                grp_data = [df[df[grp_col] == g][v_col].dropna() for g in unique_grps]
                f_stat, p_val = stats.f_oneway(*grp_data)
                results["anova"] = {
                    "f_stat": float(f_stat), 
                    "p_value": float(p_val), 
                    "significant": bool(p_val < 0.05)
                }
    except: pass

    # 6. RELATIONSHIP: CORRELATION MATRIX (Pearson r)
    if len(num_cols) > 1:
        results["correlation"] = df[num_cols].corr().to_dict()

    # 7. PREDICTIVE: LINEAR REGRESSION
    if len(num_cols) >= 2:
        try:
            # Predict the last numeric column based on the first
            slope, intercept, r, p, std = stats.linregress(df[num_cols[0]].dropna(), df[num_cols[-1]].dropna())
            results["regression"] = {
                "r_squared": float(r**2), 
                "p_value": float(p), 
                "slope": float(slope), 
                "intercept": float(intercept)
            }
        except: pass

    # 8. PSYCHOMETRIC: RELIABILITY (Cronbach's Alpha)
    if len(num_cols) >= 3:
        try:
            alpha = pg.cronbach_alpha(data=df[num_cols].dropna())[0]
            results["reliability"] = {"alpha": float(alpha)}
        except: pass

    # 9. ECOLOGICAL: SHANNON-WIENER DIVERSITY (H')
    try:
        abund_col = next((c for c in df.columns if 'count' in c.lower() or 'abundance' in c.lower()), None)
        if abund_col:
            counts = df[abund_col].dropna().values
            p = counts / np.sum(counts)
            h = -np.sum(p[p > 0] * np.log(p[p > 0]))
            results["ecological"] = {
                "shannon_h": float(h), 
                "evenness": float(h/np.log(len(counts)) if len(counts) > 1 else 0)
            }
    except: pass

    return results

@app.post("/spss/analyze")
async def analyze(file: UploadFile = File(...), options: str = Form("{}")):
    contents = await file.read()
    try:
        opt_dict = json.loads(options) if options else {}
        # Handle file types: Word, CSV, Excel
        if file.filename.endswith(".docx"):
            doc = Document(io.BytesIO(contents))
            if not doc.tables: return {"error": "No table found in Word document."}
            data = [[cell.text for cell in row.cells] for row in doc.tables[0].rows]
            df = pd.DataFrame(data[1:], columns=data[0])
        elif file.filename.endswith(".csv"):
            df = pd.read_csv(io.BytesIO(contents))
        else:
            df = pd.read_excel(io.BytesIO(contents))
            
        return run_ultimate_computation_engine(df, opt_dict)
    except Exception as e:
        return {"error": f"Neural Engine Error: {str(e)}"}
